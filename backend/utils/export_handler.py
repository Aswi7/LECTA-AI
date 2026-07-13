import os
import sys
import logging
from datetime import datetime, UTC

# Add root directory to sys.path to allow importing 'config'
root_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
if root_path not in sys.path:
    sys.path.insert(0, root_path)

from typing import Dict, Any

# PDF Imports
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor, lightgrey, white, black
from reportlab.lib.units import inch

# DOCX Imports
from docx import Document
from docx.shared import Pt

# Config Import
from config import CONFIG

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Languages Mapping
LANGUAGES_MAP = {
    'en': 'English',
    'ta': 'Tamil',
    'hi': 'Hindi',
    'te': 'Telugu',
    'kn': 'Kannada',
    'ml': 'Malayalam',
    'bn': 'Bengali',
    'mr': 'Marathi',
    'gu': 'Gujarati',
    'pa': 'Punjabi',
    'ur': 'Urdu'
}

# Register Unicode fonts for ReportLab if running on Windows
def register_unicode_fonts():
    font_name = 'Helvetica'
    bold_font_name = 'Helvetica-Bold'
    
    font_path = r"C:\Windows\Fonts\Nirmala.ttc"
    if os.path.exists(font_path):
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            pdfmetrics.registerFont(TTFont('Nirmala', font_path))
            pdfmetrics.registerFont(TTFont('Nirmala-Bold', font_path))
            font_name = 'Nirmala'
            bold_font_name = 'Nirmala-Bold'
        except Exception as e:
            logger.warning(f"Error registering Nirmala font for PDF: {e}")
            
    return font_name, bold_font_name


def export_as_pdf(session_data: Dict[str, Any], output_path: str) -> str:
    """Generates a structured PDF lecture notes document."""
    logger.info(f"Generating PDF for session {session_data.get('session_id')} at {output_path}")
    
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Register and setup font
    font_name, bold_font_name = register_unicode_fonts()
    
    # Custom Styles
    styles.add(ParagraphStyle(
        name='Answer',
        parent=styles['Normal'],
        textColor=lightgrey,
        leftIndent=20
    ))
    
    # Update stylesheet to use Unicode-capable fonts
    for style_name in list(styles.byName.keys()):
        style = styles[style_name]
        if hasattr(style, 'fontName'):
            orig_font = style.fontName or ""
            if "Bold" in orig_font:
                style.fontName = bold_font_name
            else:
                style.fontName = font_name
            
    story = []
    
    # Header: Colored Bar and Title
    header_table = Table([[Paragraph("<b>Lecture Notes</b>", ParagraphStyle(name='TitleStyle', parent=styles['Title'], textColor=white))]],
                        colWidths=[letter[0] - 1.5*inch])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), HexColor('#4A90E2')),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 12))
    
    # Metadata
    filename = session_data.get('filename', 'Unknown')
    gen_date = datetime.now(UTC).strftime("%Y-%m-%d %H:%M:%S")
    story.append(Paragraph(f"<b>File:</b> {filename}", styles['Normal']))
    story.append(Paragraph(f"<b>Generated:</b> {gen_date}", styles['Normal']))
    
    # Languages
    detected_lang = session_data.get('language', {}).get('name', 'Not available')
    target_lang_code = session_data.get('target_language')
    target_lang = LANGUAGES_MAP.get(target_lang_code, target_lang_code).upper() if target_lang_code else 'Not available'
    story.append(Paragraph(f"<b>Detected Language:</b> {detected_lang}", styles['Normal']))
    story.append(Paragraph(f"<b>Target Language:</b> {target_lang}", styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Section: Summary
    story.append(Paragraph("Summary", styles['Heading1']))
    summary_text = session_data.get('summary') or "Not available"
    story.append(Paragraph(summary_text, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Section: Key Points
    story.append(Paragraph("Key Points", styles['Heading1']))
    bullet_notes = session_data.get('bullet_notes')
    if bullet_notes:
        items = [ListItem(Paragraph(note.replace('• ', ''), styles['Normal'])) for note in bullet_notes]
        story.append(ListFlowable(items, bulletType='bullet'))
    else:
        story.append(Paragraph("Not available", styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Section: Keywords
    story.append(Paragraph("Keywords", styles['Heading1']))
    keywords = session_data.get('concepts', {}).get('keywords', [])
    if keywords:
        kw_strings = [f"{kw['keyword']} [{kw['score']:.2f}]" for kw in keywords]
        story.append(Paragraph(", ".join(kw_strings), styles['Normal']))
    else:
        story.append(Paragraph("Not available", styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Section: Translation
    story.append(Paragraph(f"Translation ({target_lang})", styles['Heading1']))
    translation = session_data.get('translated_content') or "Not available"
    story.append(Paragraph(translation, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Section: Exam Questions
    story.append(Paragraph("Exam Questions", styles['Heading1']))
    questions = session_data.get('questions', [])
    if questions:
        q_types = {'definition': 'Definitions', 'fill_blank': 'Fill in the Blanks', 'true_false': 'True or False'}
        for q_type, q_label in q_types.items():
            type_qs = [q for q in questions if q['type'] == q_type]
            if type_qs:
                story.append(Paragraph(f"<u>{q_label}</u>", styles['Heading2']))
                for q in type_qs:
                    story.append(Paragraph(f"Q: {q['question']}", styles['Normal']))
                    story.append(Paragraph(f"A: {q['answer']}", styles['Answer']))
                    story.append(Spacer(1, 6))
    else:
        story.append(Paragraph("Not available", styles['Normal']))
    
    # Footer and Page Numbering Helper
    def add_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 9)
        session_id = session_data.get('session_id', 'N/A')
        footer_text = f"Session ID: {session_id} | Generated on: {gen_date}"
        canvas.drawString(inch, 0.5 * inch, footer_text)
        canvas.drawRightString(letter[0] - inch, 0.5 * inch, f"Page {doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
    return output_path

def export_as_docx(session_data: Dict[str, Any], output_path: str) -> str:
    """Generates a structured Word document (DOCX)."""
    logger.info(f"Generating DOCX for session {session_data.get('session_id')} at {output_path}")
    
    doc = Document()
    doc.add_heading("Lecture Notes", 0)
    
    filename = session_data.get('filename', 'Unknown')
    gen_date = datetime.now(UTC).strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"File: {filename}")
    doc.add_paragraph(f"Generated: {gen_date}")
    
    detected_lang = session_data.get('language', {}).get('name', 'Not available')
    target_lang_code = session_data.get('target_language')
    target_lang = LANGUAGES_MAP.get(target_lang_code, target_lang_code).upper() if target_lang_code else 'Not available'
    doc.add_paragraph(f"Detected Language: {detected_lang}")
    doc.add_paragraph(f"Target Language: {target_lang}")
    
    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(session_data.get('summary') or "Not available")
    
    # Key Points
    doc.add_heading("Key Points", level=1)
    bullet_notes = session_data.get('bullet_notes')
    if bullet_notes:
        for note in bullet_notes:
            doc.add_paragraph(note.replace('• ', ''), style='List Bullet')
    else:
        doc.add_paragraph("Not available")
        
    # Keywords
    doc.add_heading("Keywords", level=1)
    keywords = session_data.get('concepts', {}).get('keywords', [])
    if keywords:
        kw_text = ", ".join([f"{kw['keyword']} ({kw['score']:.2f})" for kw in keywords])
        doc.add_paragraph(kw_text)
    else:
        doc.add_paragraph("Not available")
        
    # Translation
    doc.add_heading(f"Translation ({target_lang})", level=1)
    doc.add_paragraph(session_data.get('translated_content') or "Not available")
    
    # Exam Questions
    doc.add_heading("Exam Questions", level=1)
    questions = session_data.get('questions', [])
    if questions:
        q_types = {'definition': 'Definitions', 'fill_blank': 'Fill in the Blanks', 'true_false': 'True or False'}
        for q_type, q_label in q_types.items():
            type_qs = [q for q in questions if q['type'] == q_type]
            if type_qs:
                doc.add_heading(q_label, level=2)
                for q in type_qs:
                    doc.add_paragraph(f"Q: {q['question']}")
                    ans = doc.add_paragraph(f"A: {q['answer']}")
                    ans.italic = True
    else:
        doc.add_paragraph("Not available")
        
    doc.save(output_path)
    return output_path

def export_as_txt(session_data: Dict[str, Any], output_path: str) -> str:
    """Generates a structured plain text file (TXT)."""
    logger.info(f"Generating TXT for session {session_data.get('session_id')} at {output_path}")
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("=== LECTURE NOTES ===\n")
        f.write(f"File: {session_data.get('filename', 'Unknown')}\n")
        f.write(f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Session ID: {session_data.get('session_id', 'N/A')}\n\n")
        
        f.write("=== SUMMARY ===\n")
        f.write((session_data.get('summary') or "Not available") + "\n\n")
        
        f.write("=== KEY POINTS ===\n")
        bullet_notes = session_data.get('bullet_notes')
        if bullet_notes:
            for note in bullet_notes:
                f.write(f"- {note.replace('• ', '')}\n")
        else:
            f.write("Not available\n")
        f.write("\n")
        
        f.write("=== KEYWORDS ===\n")
        keywords = session_data.get('concepts', {}).get('keywords', [])
        if keywords:
            kw_text = ", ".join([f"{kw['keyword']} [{kw['score']:.2f}]" for kw in keywords])
            f.write(kw_text + "\n")
        else:
            f.write("Not available\n")
        f.write("\n")
        
        f.write("=== TRANSLATION ===\n")
        f.write((session_data.get('translated_content') or "Not available") + "\n\n")
        
        f.write("=== EXAM QUESTIONS ===\n")
        questions = session_data.get('questions', [])
        if questions:
            q_types = {'definition': 'Definitions', 'fill_blank': 'Fill in the Blanks', 'true_false': 'True or False'}
            for q_type, q_label in q_types.items():
                type_qs = [q for q in questions if q['type'] == q_type]
                if type_qs:
                    f.write(f"[{q_label}]\n")
                    for q in type_qs:
                        f.write(f"Q: {q['question']}\n")
                        f.write(f"A: {q['answer']}\n")
                    f.write("\n")
        else:
            f.write("Not available\n")
            
    return output_path

def generate_export(session_id: str, format: str, session_data: Dict[str, Any], exports_dir: str) -> str:
    """Routes the export request to the appropriate format generator."""
    valid_formats = ["pdf", "docx", "txt"]
    if format.lower() not in valid_formats:
        raise ValueError(f"Unsupported export format: {format}. Valid formats are: {', '.join(valid_formats)}")
    
    os.makedirs(exports_dir, exist_ok=True)
    filename = f"{session_id}_notes.{format.lower()}"
    output_path = os.path.join(exports_dir, filename)
    
    # If file exists, we could return it or regenerate. Requirements say "Return path".
    # I'll regenerate to ensure it's up to date with the latest session_data.
    
    if format.lower() == "pdf":
        return export_as_pdf(session_data, output_path)
    elif format.lower() == "docx":
        return export_as_docx(session_data, output_path)
    elif format.lower() == "txt":
        return export_as_txt(session_data, output_path)
    
    return output_path

def delete_exports(session_id: str):
    """Utility to delete all exports for a session (used by API)."""
    exports_dir = CONFIG.EXPORTS_FOLDER
    for fmt in ["pdf", "docx", "txt"]:
        path = os.path.join(exports_dir, f"{session_id}_notes.{fmt}")
        if os.path.exists(path):
            try:
                os.remove(path)
                logger.info(f"Deleted export: {path}")
            except Exception as e:
                logger.warning(f"Failed to delete export {path}: {e}")

# Note: The 'export_handler' singleton instance pattern from previous turn 
# is removed to follow the direct function signature requested.
